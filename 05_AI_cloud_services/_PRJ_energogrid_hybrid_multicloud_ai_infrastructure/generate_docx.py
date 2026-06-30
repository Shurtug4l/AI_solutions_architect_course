"""
Capstone modulo 05 - EnergoGrid hybrid multi-cloud AI infrastructure.

Genera il documento DOCX della consegna a partire da contenuti definiti in
questo file e dai dieci diagrammi PNG prodotti da `diagrams/00_generate_diagrams.py`.

Eseguire dalla root del progetto:

    uv run --with python-docx python generate_docx.py

Il documento risultante (PRJ_energogrid_hybrid_multicloud_ai_infrastructure.docx)
è inteso per essere caricato come unico file su Google Docs e condiviso via
link pubblico, come da consegna del modulo. Il PDF di accompagnamento può
essere prodotto con LibreOffice in modalità headless, ma non è previsto
dalla consegna.

Stile
-----
Convenzioni dalla skill docx-report-design:
  - palette near-black + amber accent (D4A017) usata con parsimonia
  - font preferito Aptos, fallback Calibri
  - body 11pt, line spacing 1.15
  - heading H1 24pt / H2 18pt / H3 14pt
  - numerazione H1 1. / H2 1.1 / H3 1.1.1
  - margini A4 2.5 cm, single column
  - caption immagine 9pt italic neutral_dk

Persona della prosa
-------------------
Voce dalla CLAUDE.md di progetto: studente Master AI Solutions Architecture
con esperienza professionale concorrente (Deloitte AI governance & risk,
cybersecurity). Tono solido tecnicamente, framing da learner, ogni scelta
non banale motivata contro alternative realistiche, analisi critica dei
risultati, prosa che non legge come AI-generated.

Lingua: italiano (modulo non ancora consegnato).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# === Style constants =====================================================

COLOR_PRIMARY    = RGBColor(0x1A, 0x1A, 0x1A)  # near-black
COLOR_TEXT       = RGBColor(0x33, 0x33, 0x33)  # body
COLOR_NEUTRAL_DK = RGBColor(0x6B, 0x6B, 0x6B)  # captions, secondary
COLOR_NEUTRAL_MD = RGBColor(0xC0, 0xC0, 0xC0)  # borders
COLOR_NEUTRAL_LT = "F0F0F0"                    # alternate row fill (hex no prefix)
COLOR_ACCENT     = RGBColor(0xD4, 0xA0, 0x17)  # amber
COLOR_ACCENT_HEX = "D4A017"
COLOR_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

FONT_FAMILY = "Aptos"
FONT_FALLBACK = "Calibri"

OUTPUT_FILE = "PRJ_energogrid_hybrid_multicloud_ai_infrastructure.docx"

# Submission date, pinned (was date.today(): fixed to avoid drift across rebuilds).
DOC_DATE = "2026-07-01"

# Module relative path to the diagrams.
DIAGRAMS_DIR = Path(__file__).parent / "diagrams"


# === Low-level helpers ===================================================


def _set_run_font(run, name=FONT_FAMILY, size=11, bold=False,
                  color=COLOR_TEXT, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    # East Asian font fallback to avoid Calibri leaking in for non-ASCII
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def _shade_cell(cell, fill_hex):
    """Apply a solid background fill to a table cell. Hex without '#'."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_cell_borders(cell, color="C0C0C0", size_pt=4):
    """Set thin grey borders on all four sides of a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size_pt))
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _add_horizontal_rule(doc, color_hex=COLOR_ACCENT_HEX, size_pt=12):
    """Insert a horizontal accent rule (used on the cover)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_pt))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


# === Document setup ======================================================


def _configure_page(doc):
    """A4 portrait, 2.5 cm margins on all sides."""
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)


def _configure_styles(doc):
    """Apply consistent typography to the built-in Word styles."""
    styles = doc.styles

    # Normal (body)
    normal = styles["Normal"]
    normal.font.name = FONT_FAMILY
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLOR_TEXT
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.1

    # Headings
    for level, size in [(1, 24), (2, 18), (3, 14), (4, 12)]:
        h = styles[f"Heading {level}"]
        h.font.name = FONT_FAMILY
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = COLOR_PRIMARY
        h.paragraph_format.space_before = Pt(16 if level == 1 else 10)
        h.paragraph_format.space_after = Pt(8 if level == 1 else 5)
        h.paragraph_format.keep_with_next = True


def _add_header_footer(doc, header_text):
    """Page header (brand line) and footer (page number)."""
    section = doc.sections[0]

    # Header
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run(header_text)
    _set_run_font(r, size=9, color=COLOR_NEUTRAL_DK, italic=True)

    # Footer with page number (PAGE field)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("pag. ")
    _set_run_font(r1, size=9, color=COLOR_NEUTRAL_DK)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_instr = OxmlElement("w:instrText")
    fld_instr.set(qn("xml:space"), "preserve")
    fld_instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r2 = fp.add_run()
    _set_run_font(r2, size=9, color=COLOR_NEUTRAL_DK)
    r2._r.append(fld_begin)
    r2._r.append(fld_instr)
    r2._r.append(fld_end)


# === Content helpers =====================================================


def add_heading(doc, text, level=1, numbered_prefix=None):
    p = doc.add_paragraph(style=f"Heading {level}")
    if numbered_prefix:
        r = p.add_run(f"{numbered_prefix}  ")
        _set_run_font(r, size=24 if level == 1 else 18 if level == 2 else 14,
                      bold=True, color=COLOR_ACCENT)
    r = p.add_run(text)
    _set_run_font(r, size=24 if level == 1 else 18 if level == 2 else 14,
                  bold=True, color=COLOR_PRIMARY)
    return p


def add_paragraph(doc, text, italic=False, bold=False, size=11,
                  color=COLOR_TEXT, alignment=None):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    r = p.add_run(text)
    _set_run_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_bullets(doc, items, indent=0.0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        r = p.add_run(item)
        _set_run_font(r, size=11, color=COLOR_TEXT)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(text)
    _set_run_font(r, size=9, italic=True, color=COLOR_NEUTRAL_DK)


def add_image(doc, filename, caption=None, width_cm=15.5):
    img_path = DIAGRAMS_DIR / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))
    if caption:
        add_caption(doc, caption)


def add_table(doc, headers, rows, col_widths_cm=None,
              first_col_bold=False, header_fill=COLOR_ACCENT_HEX,
              header_color=COLOR_WHITE):
    """A clean grey-bordered table with an accent header row.

    Args:
        headers: list of column header strings.
        rows: list of row lists, each a list of cell strings.
        col_widths_cm: optional list of widths in cm.
        first_col_bold: render the first column bold (label column).
        header_fill: hex fill for the header row.
        header_color: text color for header row.
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        _shade_cell(cell, header_fill)
        _set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        _set_run_font(r, size=10.5, bold=True, color=header_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Rows
    for i, row in enumerate(rows):
        fill = COLOR_NEUTRAL_LT if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            _shade_cell(cell, fill)
            _set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            bold = first_col_bold and j == 0
            _set_run_font(r, size=10, bold=bold, color=COLOR_TEXT)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Spacing after the table
    doc.add_paragraph()


def add_callout(doc, label, text):
    """A single-row two-cell pseudo-callout: amber tag + body."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Cm(3.0)
    table.columns[1].width = Cm(13.0)

    tag_cell = table.rows[0].cells[0]
    body_cell = table.rows[0].cells[1]

    _shade_cell(tag_cell, COLOR_ACCENT_HEX)
    _set_cell_borders(tag_cell, color=COLOR_ACCENT_HEX, size_pt=6)
    _shade_cell(body_cell, "FFFFFF")
    _set_cell_borders(body_cell, color="E5E5E5", size_pt=4)

    tag_cell.text = ""
    p = tag_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label.upper())
    _set_run_font(r, size=9, bold=True, color=COLOR_WHITE)
    tag_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    body_cell.text = ""
    p = body_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    _set_run_font(r, size=10.5, italic=True, color=COLOR_TEXT)
    body_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()


# === Cover page ==========================================================


def build_cover(doc):
    """Cover page: brand, title, subtitle, author, date, accent rule."""
    # Top spacing
    for _ in range(2):
        doc.add_paragraph()

    # Brand / context line
    p = doc.add_paragraph()
    r = p.add_run("MASTER  ·  AI SOLUTIONS ARCHITECT  ·  MODULO 05  ·  AI CLOUD SERVICES")
    _set_run_font(r, size=10, bold=True, color=COLOR_ACCENT)
    p.paragraph_format.space_after = Pt(4)

    _add_horizontal_rule(doc, color_hex=COLOR_ACCENT_HEX, size_pt=18)

    # Spacer
    for _ in range(2):
        doc.add_paragraph()

    # Main title (large)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Progettazione di un'infrastruttura\nAI multi-cloud")
    _set_run_font(r, size=30, bold=True, color=COLOR_PRIMARY)

    # Sub-title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Caso EnergoGrid S.p.A.")
    _set_run_font(r, size=18, bold=False, color=COLOR_NEUTRAL_DK)

    # Tagline
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run(
        "Strategia architetturale ad alto livello per la migrazione e l'operatività "
        "di workload AI verso una piattaforma ibrida multi-cloud, applicata al "
        "forecasting energetico, alla manutenzione predittiva degli asset e "
        "all'ottimizzazione del bilanciamento rete-produzione."
    )
    _set_run_font(r, size=12, italic=True, color=COLOR_TEXT)

    # Push author block toward the bottom of the page
    for _ in range(8):
        doc.add_paragraph()

    _add_horizontal_rule(doc, color_hex="E5E5E5", size_pt=6)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Autore  ")
    _set_run_font(r, size=10, bold=True, color=COLOR_NEUTRAL_DK)
    r = p.add_run("Simone La Porta")
    _set_run_font(r, size=10, color=COLOR_PRIMARY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Data  ")
    _set_run_font(r, size=10, bold=True, color=COLOR_NEUTRAL_DK)
    r = p.add_run(DOC_DATE)
    _set_run_font(r, size=10, color=COLOR_PRIMARY)

    p = doc.add_paragraph()
    r = p.add_run("Capstone di fine modulo  ·  Master ProfessionAI")
    _set_run_font(r, size=10, italic=True, color=COLOR_NEUTRAL_DK)

    _add_page_break(doc)


# === Table of contents (manual, static) ==================================


def build_toc(doc):
    add_heading(doc, "Indice", level=1)

    entries = [
        ("Executive summary",                                                   "3"),
        ("1.  Contesto e caso d'uso EnergoGrid",                                "4"),
        ("    1.1  Profilo dell'azienda",                                       "4"),
        ("    1.2  Tre scenari predittivi",                                     "4"),
        ("    1.3  Requisiti funzionali",                                       "5"),
        ("    1.4  Requisiti non funzionali",                                   "6"),
        ("2.  Catalogo dei servizi considerati",                                "7"),
        ("    2.1  PaaS AWS",                                                   "7"),
        ("    2.2  PaaS Microsoft Azure",                                       "8"),
        ("    2.3  PaaS Google Cloud",                                          "9"),
        ("    2.4  Open source e componenti on-prem",                          "10"),
        ("3.  Criteri di valutazione comparativa",                             "11"),
        ("4.  Analisi comparativa",                                            "13"),
        ("    4.1  Mappa di sintesi criteri x cluster",                        "13"),
        ("    4.2  Forze e debolezze per cluster",                             "14"),
        ("5.  Strategia architetturale concettuale",                           "17"),
        ("    5.1  Principi guida",                                            "17"),
        ("    5.2  Componenti logici della piattaforma",                       "18"),
        ("    5.3  Pattern di distribuzione dei workload",                     "20"),
        ("    5.4  Modello di responsabilità multi-cloud",                    "21"),
        ("6.  Impatto su processi e costi",                                    "22"),
        ("    6.1  Impatto organizzativo",                                     "22"),
        ("    6.2  TCO concettuale a 5 anni",                                  "23"),
        ("7.  Sicurezza, compliance e data governance",                        "25"),
        ("    7.1  Quadro normativo applicabile",                              "25"),
        ("    7.2  Controlli per dominio",                                     "26"),
        ("    7.3  Data governance e lineage",                                 "27"),
        ("    7.4  Checklist di alto livello",                                 "28"),
        ("8.  Piano di migrazione e roadmap",                                  "29"),
        ("    8.1  Strategia incrementale a tre wave",                         "29"),
        ("    8.2  Milestone e gate strategici",                               "30"),
        ("    8.3  Rischi principali e mitigazioni",                           "31"),
        ("9.  KPI di progetto e misurazione",                                  "33"),
        ("10. Conclusioni e raccomandazioni strategiche",                      "35"),
        ("Allegato A  ·  Glossario",                                           "37"),
        ("Allegato B  ·  Riferimenti tecnici",                                 "38"),
        ("Allegato C  ·  Riferimenti normativi",                               "39"),
    ]

    # Render as a borderless two-column table for clean alignment.
    table = doc.add_table(rows=len(entries), cols=2)
    table.autofit = False
    table.columns[0].width = Cm(13.5)
    table.columns[1].width = Cm(2.0)

    for i, (label, page) in enumerate(entries):
        c1 = table.rows[i].cells[0]
        c2 = table.rows[i].cells[1]
        c1.text = ""
        c2.text = ""
        # Section labels in bold, sub-sections normal
        is_section = label and not label.startswith("    ") and not label.startswith("Allegato")
        p = c1.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label)
        _set_run_font(r, size=11, bold=is_section,
                      color=COLOR_PRIMARY if is_section else COLOR_TEXT)
        p = c2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(page)
        _set_run_font(r, size=10, color=COLOR_NEUTRAL_DK)

    add_caption(
        doc,
        "Indice statico. La numerazione delle pagine è indicativa: il "
        "documento finale viene caricato su Google Docs, che ricalcola "
        "automaticamente la paginazione in base al rendering del browser."
    )

    _add_page_break(doc)


# === Executive summary ===================================================


def build_executive_summary(doc):
    add_heading(doc, "Executive summary", level=1)

    add_paragraph(
        doc,
        "EnergoGrid S.p.A. è un operatore integrato del settore energetico "
        "che gestisce produzione tradizionale e rinnovabile, sistemi di "
        "stoccaggio e una rete di distribuzione intelligente. Le tre "
        "iniziative AI prioritarie - forecasting della domanda, manutenzione "
        "predittiva degli asset e ottimizzazione del bilanciamento - hanno "
        "ricaduta diretta su affidabilità della rete, costo del kilowattora "
        "consegnato e capacità di assorbire la variabilità delle "
        "rinnovabili. Tutte e tre condividono la stessa esigenza "
        "infrastrutturale: una piattaforma AI in grado di trattare dati di "
        "telemetria con vincoli stringenti di sovranità, di sostenere il "
        "ciclo di vita dei modelli su orizzonti pluriennali e di esporre "
        "previsioni sia in tempo reale al dispacciamento sia in batch a "
        "manutenzione e trading."
    )

    add_paragraph(
        doc,
        "Il documento argomenta che la risposta corretta non è "
        "single-cloud, né tantomeno una replica completa multi-cloud, ma "
        "una architettura ibrida che mantiene on-prem il perimetro a maggior "
        "sensibilità e classifica i workload restanti su tre cloud "
        "pubblici secondo un criterio best-of-breed: AWS come spina dorsale "
        "ML, Azure per integrazione enterprise e LLM regolati, GCP per "
        "analytics su grandi volumi. Open source (Docker, Kubernetes, "
        "MLflow, Hugging Face) è la lingua franca che tiene insieme i tre "
        "ambienti pubblici con il perimetro privato, mitigando il lock-in "
        "senza costruire in casa un layer di astrazione proprietario."
    )

    add_paragraph(
        doc,
        "La proiezione concettuale a 5 anni indica un risparmio cumulato "
        "intorno al 20% rispetto a una baseline on-prem only, principalmente "
        "per la sostituzione del refresh hardware periodico con consumo "
        "elastico cloud sui picchi di training. Il taglio dei costi, però, "
        "è il beneficio minore: pesano di più la flessibilità di provare "
        "nuovi modelli senza acquistare GPU e l'allineamento con i requisiti "
        "regolatori in evoluzione (NIS2 entrata in vigore nel 2024, EU AI "
        "Act, ISO 27019 per il settore energia). La roadmap raccomandata è "
        "a tre wave su 36 mesi, con gate decisionali a M9, M21 e M36 e "
        "milestone misurabili sui tre casi d'uso."
    )

    add_callout(
        doc,
        "raccomandazione",
        "Adottare un'architettura ibrida multi-cloud guidata da workload "
        "(non da provider). Single-cloud sui workload dove il data gravity "
        "lo impone, OSS on-prem dove la sensibilità del dato e la latenza "
        "lo richiedono. Evitare il full multi-cloud replicato: il costo "
        "operativo non è giustificato dai benefici per un operatore di "
        "questa scala."
    )

    add_paragraph(
        doc,
        "Le sezioni successive sviluppano la motivazione di ciascuna scelta "
        "rispetto alle alternative considerate, presentano l'analisi "
        "comparativa sui criteri richiesti dalla consegna, descrivono la "
        "strategia architetturale, l'impatto economico e organizzativo, gli "
        "aspetti di sicurezza e compliance, la roadmap a wave, i KPI di "
        "successo e le raccomandazioni di chiusura."
    )

    _add_page_break(doc)


# === Section 1 - Context and use case ====================================


def build_section_1(doc):
    add_heading(doc, "Contesto e caso d'uso EnergoGrid",
                level=1, numbered_prefix="1.")

    # 1.1
    add_heading(doc, "Profilo dell'azienda", level=2, numbered_prefix="1.1")
    add_paragraph(
        doc,
        "EnergoGrid S.p.A. opera lungo l'intera catena del valore "
        "elettrico: produzione (centrali tradizionali termoelettriche e "
        "parchi di rinnovabili distribuiti), stoccaggio, trasporto e "
        "distribuzione su rete di media e bassa tensione. Lo scenario "
        "operativo è quello tipico di un'utility europea contemporanea: "
        "una quota crescente di generazione intermittente, una clientela "
        "che diventa anche produttore (prosumer), una pressione regolatoria "
        "in aumento sull'affidabilità del servizio e una rete di asset "
        "fisici eterogenea per età, tecnologia e fornitore."
    )
    add_paragraph(
        doc,
        "Il punto di partenza tecnologico include sistemi SCADA tradizionali "
        "per il controllo della rete, una infrastruttura di telemetria IoT "
        "con copertura non uniforme (smart meter di nuova generazione "
        "affiancati a contatori legacy), data warehouse aziendali su "
        "tecnologia on-prem e una recente sperimentazione cloud su AWS per "
        "carichi non critici. Il bagaglio AI esistente è limitato a "
        "modelli statistici di forecasting di vecchia generazione gestiti "
        "in modalità poco industrializzata."
    )

    # 1.2
    add_heading(doc, "Tre scenari predittivi", level=2, numbered_prefix="1.2")
    add_paragraph(
        doc,
        "La consegna individua tre casi d'uso AI con priorità alta per "
        "EnergoGrid. Sono trattati come un unico programma piuttosto che "
        "come tre progetti scollegati, perché condividono ingestione, "
        "feature store, governance e una larga parte dei modelli "
        "infrastrutturali."
    )

    add_table(
        doc,
        headers=["Caso d'uso", "Output operativo", "Orizzonte", "Frequenza"],
        rows=[
            ["Forecasting domanda energetica",
             "Curva di carico per nodo di rete",
             "T+1h a T+72h",
             "ogni 15 min"],
            ["Manutenzione predittiva asset",
             "Probabilità di guasto per inverter/trasformatore",
             "T+24h a T+30g",
             "ogni 1h"],
            ["Ottimizzazione bilanciamento e storage",
             "Set point di dispacciamento e cariche batteria",
             "T+15min a T+24h",
             "ogni 15 min"],
        ],
        col_widths_cm=[5.5, 5.5, 2.5, 2.5],
        first_col_bold=True,
    )

    add_paragraph(
        doc,
        "Forecasting e ottimizzazione di bilanciamento si avvalgono delle "
        "stesse features (consumi storici, meteo, prezzo zonale, stato "
        "asset); la manutenzione predittiva attinge invece dalla "
        "telemetria di basso livello degli inverter e dei trasformatori "
        "(temperature, vibrazioni, profilo di carico). Trattarli come uno "
        "stesso programma significa riusare la pipeline di ingestione e il "
        "feature store, e ridurre il costo marginale di portare in "
        "produzione il secondo e il terzo modello rispetto al primo."
    )

    # 1.3
    add_heading(doc, "Requisiti funzionali", level=2, numbered_prefix="1.3")
    add_paragraph(
        doc,
        "I requisiti funzionali sono espressi al livello di programma e "
        "non come specifica di singolo modello, coerentemente con il "
        "perimetro strategico della consegna."
    )

    add_table(
        doc,
        headers=["ID", "Requisito funzionale"],
        rows=[
            ["F1", "Ingestione continua di telemetria asset da SCADA, gateway IoT industriali e smart meter, con bufferizzazione locale a tolleranza di interruzioni di rete fino a 24 ore."],
            ["F2", "Pipeline di feature engineering condivisa fra i tre casi d'uso, con feature store come unica fonte di verità sia in training sia in serving."],
            ["F3", "Catalogo modelli versionato con stage di promozione (development, staging, production, archived) e tracciatura completa delle metriche di valutazione."],
            ["F4", "Endpoint di inferenza online per dispacciamento e ottimizzazione bilanciamento, con latenza p95 sotto 500 ms a partire dalla richiesta interna."],
            ["F5", "Batch scoring giornaliero per manutenzione predittiva, con output disponibili al sistema asset management entro le 06:00 locali."],
            ["F6", "Inferenza locale a bordo substation per anomaly detection a bassa latenza (target sotto 50 ms) sui parametri elettrici."],
            ["F7", "Monitoraggio di drift sulla distribuzione delle feature di ingresso e sulle metriche di qualità dei modelli, con trigger automatici di retraining."],
            ["F8", "Cruscotti decisionali per dispatcher e responsabili manutenzione, con accesso amministrato tramite il sistema di identità aziendale."],
            ["F9", "Esportazione audit-ready di evidenze su data lineage, decisioni del modello e accessi, in formato adatto a ispezione regolatoria."],
        ],
        col_widths_cm=[1.5, 14.5],
        first_col_bold=True,
    )

    # 1.4
    add_heading(doc, "Requisiti non funzionali", level=2, numbered_prefix="1.4")
    add_paragraph(
        doc,
        "I requisiti non funzionali sono i vincoli che pesano sulle scelte "
        "architetturali più di qualunque preferenza tecnologica. "
        "In un operatore di rete elettrica la priorità non è la novità "
        "del modello, è che il sistema operi 24 ore su 24 e non sia un "
        "vettore di compromissione dell'infrastruttura critica."
    )

    add_table(
        doc,
        headers=["ID", "Categoria", "Requisito non funzionale"],
        rows=[
            ["NF1", "Disponibilità",
             "RTO sotto 4 ore e RPO sotto 1 ora per le pipeline che alimentano dispatch e manutenzione."],
            ["NF2", "Latenza",
             "p95 sotto 500 ms per inferenza online, p95 sotto 50 ms per anomaly detection a substation."],
            ["NF3", "Sovranità del dato",
             "Telemetria asset e dati operativi devono risiedere in territorio italiano o EU; nessuna esportazione verso jurisdiction non adeguate ai sensi del GDPR."],
            ["NF4", "Sicurezza",
             "Conformità a NIS2 per operatori essenziali, ISO 27001/27019, IEC 62443 per la parte OT, separazione netta IT-OT a livello di rete e di identità."],
            ["NF5", "Interoperabilità",
             "Integrazione con SCADA esistenti via protocolli industriali (IEC 60870-5-104, OPC UA) senza riprogettazione degli stessi."],
            ["NF6", "Portabilità",
             "Modelli e pipeline devono poter migrare fra cloud provider con un effort di settimane, non di mesi; il vendor lock-in è inaccettabile come default architetturale."],
            ["NF7", "Auditabilità",
             "Tracciatura completa di chi ha addestrato cosa con quali dati, conservata per almeno 5 anni in conformità ai requisiti di audit interno e regolatorio."],
            ["NF8", "Sostenibilità operativa",
             "Il team interno cresce in modo graduale; la piattaforma deve essere operabile da un nucleo iniziale di 6-10 persone, non da un'organizzazione di 50."],
            ["NF9", "Costo",
             "TCO a 5 anni inferiore al 90% della baseline on-prem only; cost shape prevedibile e budgettabile per ciclo finanziario annuale."],
        ],
        col_widths_cm=[1.2, 3.5, 11.3],
        first_col_bold=True,
    )

    add_callout(
        doc,
        "punto critico",
        "I requisiti NF2, NF3 e NF4 escludono un'architettura puramente "
        "cloud per la parte vicina agli asset fisici. NF6 e NF9 escludono "
        "una soluzione single-vendor PaaS spinta. L'intersezione dei due "
        "vincoli porta in modo naturale a un disegno ibrido: questa non è "
        "una preferenza estetica, è l'unica scelta che soddisfi i "
        "requisiti minimi senza eccezioni esplicite."
    )

    _add_page_break(doc)


# === Section 2 - Catalogue ===============================================


def build_section_2(doc):
    add_heading(doc, "Catalogo dei servizi considerati",
                level=1, numbered_prefix="2.")

    add_paragraph(
        doc,
        "Il catalogo include le piattaforme PaaS dei tre principali "
        "hyperscaler e il sottoinsieme open source rilevante per il caso "
        "EnergoGrid. La descrizione si mantiene di proposito a livello alto: il "
        "documento non valuta specifiche tecniche di servizio, ma il "
        "contributo di ciascun cluster alla strategia. La selezione "
        "all'interno di ciascun PaaS si limita ai servizi pertinenti al "
        "ciclo di vita ML, all'ingestione di telemetria industriale e alla "
        "data foundation."
    )

    # 2.1 AWS
    add_heading(doc, "PaaS AWS", level=2, numbered_prefix="2.1")
    add_paragraph(
        doc,
        "AWS è il provider con il catalogo più ampio e con la maturità "
        "operativa più lunga. Per il programma EnergoGrid è candidato a "
        "ruolo di spina dorsale ML: SageMaker copre l'intera filiera di "
        "training, registry, deployment e drift monitoring, mentre IoT "
        "Core fornisce un canale di ingestione MQTT gestito con device "
        "registry e shadow state per la parte di gateway industriali. La "
        "scelta non implica esclusività: i workload sensibili restano on-prem "
        "e altri cloud entrano in scena dove portano vantaggi specifici."
    )
    add_table(
        doc,
        headers=["Servizio AWS", "Ruolo nella piattaforma"],
        rows=[
            ["Amazon S3", "Data lake unico per dataset di training, modelli, log e artefatti, con bucket dedicati per perimetro di classificazione."],
            ["Amazon SageMaker", "Piattaforma end-to-end per training, hyperparameter tuning, model registry, endpoint REST e Model Monitor per drift detection."],
            ["Amazon Aurora", "Feature store transazionale e archivio delle previsioni con consistenza ACID."],
            ["Amazon DynamoDB", "Lookup a singola chiave a bassa latenza per cache di feature online e stato utente delle dashboard operative."],
            ["AWS IoT Core", "Ingestione MQTT da gateway industriali, device shadow, regole di routing verso Kinesis o Lambda."],
            ["Amazon CloudWatch + EventBridge", "Osservabilità di pipeline e endpoint, automazione dei trigger di retraining basati su soglie di drift."],
            ["AWS KMS + Secrets Manager", "Gestione chiavi cifratura e segreti, integrabile con HSM esterno per scenari ad alta sensibilità."],
        ],
        col_widths_cm=[5.5, 11.0],
        first_col_bold=True,
    )

    # 2.2 Azure
    add_heading(doc, "PaaS Microsoft Azure", level=2,
                numbered_prefix="2.2")
    add_paragraph(
        doc,
        "Azure ha il vantaggio strutturale dell'integrazione con la "
        "filiera Microsoft (Active Directory, Office 365, Dynamics) e il "
        "portafoglio di certificazioni più ampio per i settori regolati. "
        "Per EnergoGrid il ruolo naturale è duplice: ospitare i workload "
        "che si appoggiano all'identità aziendale e ai canali enterprise "
        "(reportistica per management, integrazione con strumenti di "
        "ticketing), e fornire l'accesso a modelli generativi (Azure "
        "OpenAI) sotto contratto con garanzie di non-training sui dati e "
        "data residency in regione EU."
    )
    add_table(
        doc,
        headers=["Servizio Azure", "Ruolo nella piattaforma"],
        rows=[
            ["Azure Blob Storage", "Storage oggetti per dataset di training secondari e archiviazione storica."],
            ["Azure Machine Learning", "Piattaforma di training, AutoML, MLOps, endpoint gestiti per i modelli che si integrano con i sistemi Microsoft di reportistica."],
            ["Azure OpenAI Service", "Accesso ai modelli GPT-4o / o-series in regione EU con garanzie contrattuali enterprise, per assistenza documentale e generazione di report."],
            ["Azure AI Foundry", "Hub per la sperimentazione di agenti e LLM customizzati, isolato dal perimetro produttivo."],
            ["Azure Cosmos DB", "Storico delle conversazioni operatore-assistente e cache geograficamente distribuita."],
            ["Azure IoT Hub", "Canale alternativo o ridondato di ingestione per gateway che vivono nell'ecosistema Microsoft."],
            ["Azure Active Directory / Entra ID", "Identità unificata, single sign-on, accesso condizionale e MFA per gli operatori."],
        ],
        col_widths_cm=[5.5, 11.0],
        first_col_bold=True,
    )

    # 2.3 GCP
    add_heading(doc, "PaaS Google Cloud", level=2, numbered_prefix="2.3")
    add_paragraph(
        doc,
        "GCP si distingue per la filosofia data-first: BigQuery è il "
        "motore di analytics serverless di riferimento sul mercato e "
        "BigQuery ML porta il training dentro il warehouse, con modelli "
        "addestrati direttamente sui dati e senza esportazioni. Vertex AI "
        "fornisce il completamento naturale verso il training custom e la "
        "gestione del lifecycle. Per "
        "EnergoGrid GCP entra in scena nella parte analitica del programma: "
        "consolidamento storico dei consumi, segmentazioni di lungo periodo, "
        "analisi pattern di rete che richiedono scansioni su anni di "
        "telemetria."
    )
    add_table(
        doc,
        headers=["Servizio GCP", "Ruolo nella piattaforma"],
        rows=[
            ["Google Cloud Storage", "Storage oggetti per snapshot del data warehouse e backup."],
            ["BigQuery", "Data warehouse analitico per consumi storici, mercato e dataset di forecasting di lungo periodo."],
            ["BigQuery ML", "Training di modelli tabellari (forecasting baseline, segmentazione clienti) direttamente in SQL."],
            ["Vertex AI", "Piattaforma di training custom, model registry, endpoint, Feature Store, monitoring di drift; alternativa o complemento a SageMaker."],
            ["Vertex AI Search", "Retrieval ibrido vettoriale e keyword per casi d'uso di RAG su documentazione tecnica e normativa."],
            ["Dataflow + Pub/Sub", "Streaming ETL e ingestione real-time per integrazione con feed di mercato e meteo ad alta frequenza."],
        ],
        col_widths_cm=[5.5, 11.0],
        first_col_bold=True,
    )

    # 2.4 OSS + on-prem
    add_heading(doc, "Open source e componenti on-prem", level=2,
                numbered_prefix="2.4")
    add_paragraph(
        doc,
        "La componente open source affianca il cloud invece di sostituirlo: "
        "fa da strato comune che rende l'architettura portabile fra cloud e "
        "verso on-prem. Tre vantaggi strutturali ne giustificano la "
        "presenza in ogni layer: portabilità dei modelli e dei container, "
        "controllo del codice per audit di terze parti, riduzione del "
        "lock-in senza la complessità di un layer di astrazione proprietario."
    )
    add_table(
        doc,
        headers=["Componente OSS / on-prem", "Ruolo nella piattaforma"],
        rows=[
            ["Docker", "Standard di packaging per ogni modello, libreria e servizio della piattaforma; abilita la portabilità fra ambienti."],
            ["Kubernetes (K8s) + K3s", "Orchestrazione cluster on-prem (K8s) e edge a substation (K3s lightweight), con stessa control plane semantica."],
            ["MLflow", "Tracking degli esperimenti e Model Registry open source, fonte unica di verità per il lifecycle dei modelli indipendentemente dal cloud in cui il training avviene."],
            ["Hugging Face Hub", "Catalogo modelli pre-addestrati (transformer per testo, time-series), riferimento per supply chain dei modelli con versioning e model card."],
            ["FastAPI", "Layer di serving Python, con OpenAPI automatico per integrazione con i sistemi consumer."],
            ["Apache Kafka", "Backbone di streaming dati per ingestione telemetria; standard di settore con ecosistema connector ampio."],
            ["MinIO", "Object storage S3-compatible on-prem per i dataset sensibili che non lasciano il perimetro."],
            ["TimescaleDB", "Time-series database PostgreSQL-compatible, idoneo per storicizzazione telemetria a media frequenza."],
            ["Prometheus + Grafana", "Stack di osservabilità end-to-end di metriche infrastrutturali, applicative e di modello."],
            ["HashiCorp Vault", "Gestione segreti e chiavi cifratura unificata fra cloud e on-prem, riduce il drift fra IAM dei singoli provider."],
        ],
        col_widths_cm=[5.5, 11.0],
        first_col_bold=True,
    )

    add_callout(
        doc,
        "principio guida",
        "Ogni servizio PaaS adottato deve essere accessibile da codice "
        "interno tramite un'astrazione che non lo nomini esplicitamente "
        "fuori dal modulo di integrazione. Così la portabilità resta "
        "un'opzione operativa, non una promessa marketing."
    )

    _add_page_break(doc)


# === Section 3 - Evaluation criteria =====================================


def build_section_3(doc):
    add_heading(doc, "Criteri di valutazione comparativa",
                level=1, numbered_prefix="3.")

    add_paragraph(
        doc,
        "La consegna richiede otto criteri di valutazione. Tutti sono "
        "strategici (non operativi) e sono interpretati come domande sul "
        "ruolo che ciascun cluster di soluzioni può coprire nella "
        "piattaforma EnergoGrid. Per ogni criterio sono indicati: la "
        "definizione operativa, i driver concreti e il metodo di confronto "
        "applicato nella sezione 4."
    )

    add_paragraph(
        doc,
        "Il metodo è qualitativo su scala 1-5, discreto per scelta. "
        "Una valutazione su scala continua avrebbe richiesto stime "
        "numeriche puntuali (per esempio euro / mese / req) che a livello "
        "concettuale sarebbero state pretese di precisione non "
        "giustificabili. La granularità 1-5 è sufficiente a discriminare "
        "i cluster e lascia esplicito il margine di errore."
    )

    add_table(
        doc,
        headers=["Scala", "Etichetta", "Significato"],
        rows=[
            ["1", "debole",         "non soddisfa il requisito senza interventi sostanziali"],
            ["2", "sotto la media", "soddisfa parzialmente, gap rilevanti"],
            ["3", "adeguato",       "soddisfa il requisito a livello base"],
            ["4", "forte",          "soddisfa il requisito con margine, pochi gap"],
            ["5", "best-in-class",  "riferimento di mercato sul criterio"],
        ],
        col_widths_cm=[1.5, 3.0, 11.5],
        first_col_bold=True,
    )

    # Criteri uno per uno
    criteri = [
        ("3.1", "Costo (TCO 3-5 anni)",
         "Costo totale di possesso su orizzonte 3 e 5 anni, includendo "
         "spese ricorrenti (run-time, storage, traffico), licenze, sforzo "
         "di operations e refresh hardware dove applicabile.",
         [
             "spesa per inferenza (per call su PaaS, per ora-GPU su OSS)",
             "spesa per storage e per egress cross-cloud",
             "licenze enterprise (osservabilità, security, supporto)",
             "ammortamento hardware e ciclo di refresh (solo on-prem)",
             "costo del personale di operations dedicato",
         ]),
        ("3.2", "Scalabilità e resilienza",
         "Capacità di scalare training e inference, gestire picchi di "
         "carico (per esempio in fase di tempesta che induce stress sulla "
         "rete) e garantire ridondanza geografica.",
         [
             "elasticità di compute e storage",
             "tempo di provisioning per scenario di picco",
             "modello di disaster recovery e RTO/RPO",
             "supporto a multi-region con failover automatico",
         ]),
        ("3.3", "Sicurezza e compliance",
         "Insieme dei controlli di sicurezza nativi e delle certificazioni "
         "che riducono lo sforzo di compliance per il settore energia.",
         [
             "controlli IAM granulari e supporto per separazione IT-OT",
             "cifratura at-rest e in-transit, gestione chiavi (KMS, HSM)",
             "isolamento multi-tenant, network segmentation, private link",
             "auditability (audit log, immutabilità, conservazione)",
             "certificazioni (ISO 27001/27019, SOC 2, ENISA, conformità NIS2, EU AI Act)",
         ]),
        ("3.4", "Portabilità e rischio di vendor lock-in",
         "Costo di migrazione e interoperabilità fra ambienti. E' una "
         "variabile strategica per un operatore di rete: la durata di "
         "esercizio dei sistemi è decennale, le scelte di provider "
         "andranno riviste più di una volta.",
         [
             "dipendenza da API proprietarie vs standard aperti",
             "formato di interscambio per modelli e dataset",
             "facilità di eseguire la stessa pipeline su provider alternativo",
             "presenza di clausole contrattuali sulla portabilità",
         ]),
        ("3.5", "Operabilità e ciclo di vita ML",
         "Supporto end-to-end al ciclo di vita: experiment tracking, "
         "model registry, deployment, monitoring di drift, retraining "
         "automatizzato.",
         [
             "experiment tracking integrato",
             "model registry con stage e approvazione",
             "deployment automatizzato a endpoint o batch",
             "monitoring di drift di input e output",
             "supporto a CI/CD/CT (continuous training)",
         ]),
        ("3.6", "Ecosistema e integrazione",
         "Disponibilità di connettori e SDK verso i sistemi che la "
         "piattaforma deve consumare o servire: SCADA, IoT industriali, "
         "ERP, asset management, CRM.",
         [
             "connettori per protocolli industriali (OPC UA, IEC 60870-5-104, Modbus)",
             "SDK e bridge verso ERP (SAP, Oracle) e CRM",
             "integrazione DevOps con repository di codice e CI",
             "ricchezza dell'ecosistema marketplace e partner",
         ]),
        ("3.7", "Supporto e comunità",
         "Disponibilità di supporto commerciale (SLA, account "
         "management, escalation path) e robustezza della comunità open "
         "source di riferimento.",
         [
             "presenza locale (Italia / EU) del supporto vendor",
             "tempo di risposta su severity 1",
             "dimensione e attività della comunità OSS per i componenti adottati",
         ]),
        ("3.8", "Time-to-market e agilità",
         "Tempo per portare in produzione una nuova capacità AI dalla "
         "decisione di business al primo rollout pilota.",
         [
             "tempo di provisioning di un nuovo workload",
             "presenza di template e accelerator pronti all'uso",
             "snellezza del processo di approvazione interno indotto dalla soluzione",
         ]),
    ]

    for num, name, desc, drivers in criteri:
        add_heading(doc, name, level=2, numbered_prefix=num)
        add_paragraph(doc, desc)
        p = doc.add_paragraph()
        r = p.add_run("Driver principali:")
        _set_run_font(r, size=11, bold=True, color=COLOR_PRIMARY)
        add_bullets(doc, drivers)

    _add_page_break(doc)


# === Section 4 - Comparative analysis ====================================


def build_section_4(doc):
    add_heading(doc, "Analisi comparativa", level=1, numbered_prefix="4.")

    # 4.1
    add_heading(doc, "Mappa di sintesi criteri x cluster",
                level=2, numbered_prefix="4.1")
    add_paragraph(
        doc,
        "La figura 4.1 sintetizza il punteggio assegnato a ciascun "
        "cluster di soluzioni sui criteri della sezione 3. La lettura "
        "raccomandata è per colonna: il profilo di ciascun cluster emerge "
        "dal pattern di forze e debolezze, non dal punteggio aggregato. "
        "Aggregare in una media unica sarebbe ingannevole, perché i "
        "criteri non hanno tutti lo stesso peso per EnergoGrid (sicurezza, "
        "portabilità e operabilità pesano più di time-to-market in un "
        "settore regolato)."
    )
    add_image(doc, "03_criteria_solutions_heatmap.png",
              caption="Figura 4.1 - Mappa qualitativa criteri x cluster di soluzioni. "
                      "Scala 1-5, lettura per colonna.")

    # 4.2
    add_heading(doc, "Forze e debolezze per cluster",
                level=2, numbered_prefix="4.2")
    add_paragraph(
        doc,
        "I cinque cluster sono valutati di seguito, ciascuno con il "
        "profilo che emerge dalla heatmap, le forze sostanziali per il "
        "caso EnergoGrid e i limiti che devono essere mitigati per renderli "
        "utilizzabili in produzione."
    )

    # AWS
    add_heading(doc, "AWS PaaS", level=3, numbered_prefix="4.2.1")
    add_paragraph(
        doc,
        "Profilo: best-in-class su scalabilità, operabilità ML, "
        "ecosistema e supporto; forte su sicurezza; adeguato su costi; "
        "sotto la media su portabilità."
    )
    add_paragraph(
        doc,
        "AWS ha la piattaforma ML più matura del mercato: SageMaker copre "
        "con un solo set di API training, registry, deployment e "
        "monitoring, IoT Core regge l'ingestione MQTT a scala industriale, "
        "e la maturità si traduce in documentazione, runbook e partner "
        "presenti in Italia. Il punto debole è la portabilità: uscire da "
        "SageMaker richiede di riscrivere la filiera di training e serving. "
        "La compliance generale è solida (ISO 27001, EU Cloud CoC), ma "
        "quella specifica per l'energia (ISO 27019, IEC 62443) va "
        "verificata puntualmente, non è garantita come pacchetto unico."
    )

    # Azure
    add_heading(doc, "Azure PaaS", level=3, numbered_prefix="4.2.2")
    add_paragraph(
        doc,
        "Profilo: best-in-class su sicurezza e compliance, "
        "scalabilità, operabilità ML, ecosistema e supporto; adeguato "
        "su costi; sotto la media su portabilità."
    )
    add_paragraph(
        doc,
        "Azure ha il portafoglio di certificazioni più ampio del settore "
        "(GDPR, ISO 27001/27017/27018, FedRAMP, ENISA Cloud Security e "
        "ISO/IEC 27019 esplicito per l'energia) e l'integrazione nativa con "
        "Active Directory e la filiera Microsoft riduce attrito su identità "
        "aziendale e reportistica per il management. Azure OpenAI Service è "
        "la via più diretta verso modelli generativi sotto vincoli "
        "enterprise. Il limite è lo stesso di AWS sulla portabilità (SDK e "
        "formati non standard de facto); l'integrazione con i protocolli "
        "industriali (OPC UA, IEC 104) non è first-class e richiede "
        "componenti aggiuntivi o partner certificati."
    )

    # GCP
    add_heading(doc, "GCP PaaS", level=3, numbered_prefix="4.2.3")
    add_paragraph(
        doc,
        "Profilo: best-in-class su scalabilità, operabilità ML e "
        "time-to-market; forte su sicurezza ed ecosistema; adeguato su "
        "portabilità e costi."
    )
    add_paragraph(
        doc,
        "GCP è il provider migliore quando il problema è dominato dai dati. "
        "BigQuery non ha equivalenti immediati: petabyte di telemetria "
        "interrogabili in SQL senza gestire cluster, e BigQuery ML addestra "
        "modelli tabellari (forecasting baseline, segmentazione di rete) "
        "direttamente sul warehouse, in minuti e senza esportazioni; Vertex "
        "AI chiude verso i modelli custom. Le debolezze sono la presenza "
        "enterprise in Italia meno diffusa rispetto ad AWS e Azure (supporto "
        "in inglese, meno partner locali) e certificazioni per l'energia "
        "presenti ma meno consolidate nei dossier regolatori."
    )

    # OSS on-prem
    add_heading(doc, "OSS on-prem", level=3, numbered_prefix="4.2.4")
    add_paragraph(
        doc,
        "Profilo: best-in-class su costi e portabilità; forte su "
        "sicurezza e supporto; adeguato su scalabilità, operabilità e "
        "ecosistema; sotto la media su time-to-market."
    )
    add_paragraph(
        doc,
        "L'OSS on-prem è la risposta corretta dove il dato non può uscire "
        "dal perimetro e la latenza deve restare sotto il centinaio di "
        "millisecondi. Lo stack di riferimento (Kubernetes + MLflow + "
        "FastAPI + Prometheus + Grafana) è standard nel settore, con "
        "comunità attive e supporto enterprise (Red Hat, SUSE, Canonical) "
        "per chi vuole garanzie contrattuali; ammortizzato l'hardware, il "
        "costo unitario di inferenza è il più basso fra le opzioni per "
        "volumi prevedibili. Il costo nascosto è il time-to-market: "
        "allestire un cluster ML on-prem maturo richiede mesi (hardware, "
        "rete, policy IAM) contro i giorni del cloud, e la scalabilità "
        "elastica resta limitata dal parco hardware, lenta sui picchi "
        "imprevisti."
    )

    # Hybrid mix
    add_heading(doc, "Hybrid mix", level=3, numbered_prefix="4.2.5")
    add_paragraph(
        doc,
        "Profilo: best-in-class su sicurezza, compliance ed "
        "ecosistema; forte sugli altri criteri; adeguato sui costi."
    )
    add_paragraph(
        doc,
        "L'hybrid mix non è un cluster a sé stante: è la "
        "combinazione strategica dei precedenti. Il suo profilo nella "
        "heatmap è figlio della scomposizione dei workload: il forecasting "
        "su volumi prevedibili va su cloud PaaS, la manutenzione "
        "predittiva sensibile resta on-prem, le componenti generative "
        "vanno su Azure OpenAI per garanzie contrattuali, le aggregazioni "
        "su BigQuery dove i dati storici lo giustificano. Ogni cluster "
        "porta la sua forza, mentre il prezzo da pagare è la complessità "
        "operativa: il rischio operativo dell'hybrid è reale e gestirlo "
        "consuma una parte sostanziale del beneficio (sezione 6)."
    )

    add_callout(
        doc,
        "lettura sintetica",
        "Nessun cluster vince da solo su tutti i criteri rilevanti per "
        "EnergoGrid. AWS e Azure perdono sulla portabilità, GCP perde "
        "sulla presenza locale, OSS on-prem perde sul time-to-market, "
        "l'hybrid perde sulla complessità operativa. La strategia "
        "raccomandata è un hybrid disciplinato, costruito attorno al "
        "principio di assegnare ogni workload al cluster dove la sua "
        "intersezione di vincoli rende l'opzione dominante."
    )

    _add_page_break(doc)


# === Section 5 - Architectural strategy ==================================


def build_section_5(doc):
    add_heading(doc, "Strategia architetturale concettuale",
                level=1, numbered_prefix="5.")

    # 5.1
    add_heading(doc, "Principi guida", level=2, numbered_prefix="5.1")
    add_paragraph(
        doc,
        "Cinque principi guidano ogni scelta architetturale a valle. "
        "Funzionano come criteri di decisione binari, applicati in modo "
        "deterministico quando emerge un'ambiguità sulla collocazione di "
        "un componente."
    )

    add_table(
        doc,
        headers=["Principio", "Applicazione operativa"],
        rows=[
            ["Edge-to-cloud per default",
             "I dati nascono all'edge, vengono aggregati on-prem e replicati selettivamente in cloud. Il percorso inverso è eccezione, non regola."],
            ["Data sovereignty inderogabile",
             "Telemetria asset e dati operativi restano in jurisdiction EU. Qualsiasi servizio PaaS che non garantisca residency esplicita è escluso."],
            ["Portabilità by design",
             "Ogni servizio PaaS è incapsulato in un'astrazione interna che ne nasconde l'origine. Sostituirlo richiede settimane, non mesi."],
            ["Open source come collante",
             "Le interfacce fra cluster (registry modelli, observability, IaC, identità) usano standard open source. Le scelte proprietarie sono ammesse solo dentro i confini di un singolo cluster."],
            ["Operability prima dell'efficienza",
             "Una architettura che il team non sa operare è un debito tecnico, non un asset. Le decisioni privilegiano l'osservabilità e la semplicità operativa anche a costo di una marginale perdita di efficienza."],
        ],
        col_widths_cm=[5.0, 11.5],
        first_col_bold=True,
    )

    # 5.2
    add_heading(doc, "Componenti logici della piattaforma",
                level=2, numbered_prefix="5.2")
    add_paragraph(
        doc,
        "La piattaforma è strutturata in quattro layer logici "
        "indipendenti dal provider che li ospita. La separazione è "
        "intenzionale: ogni layer è realizzato dal cluster di soluzioni "
        "che meglio lo serve, e cambia provider in modo controllato "
        "senza riprogettazione del resto."
    )

    add_image(doc, "01_hybrid_multicloud_architecture.png",
              caption="Figura 5.1 - Architettura concettuale ibrida multi-cloud. "
                      "Quattro layer (edge, private, public, consumers) collegati "
                      "da flussi verticali unidirezionali top-down; cross-cloud "
                      "assente per design.",
              width_cm=16.0)

    add_paragraph(
        doc,
        "Il flusso del dato attraversa i quattro layer in modo "
        "asimmetrico: la maggior parte del traffico sale dagli asset "
        "fisici verso la piattaforma analitica (telemetria, eventi, stato "
        "operativo), mentre solo decisioni e modelli aggiornati scendono "
        "in senso opposto. La figura 5.2 esplicita questa direzionalità "
        "e la classifica delle categorie di dato che percorrono la "
        "pipeline, dato che la sensibilità di ciascuna determina dove "
        "può essere processata."
    )

    add_image(doc, "02_edge_to_cloud_data_flow.png",
              caption="Figura 5.2 - Flusso dati edge-to-cloud. Sei stadi "
                      "sequenziali dalla sorgente al consumer; il loop di "
                      "retraining su drift è descritto nella sezione 7.",
              width_cm=16.0)

    add_table(
        doc,
        headers=["Layer", "Cosa fa", "Dove vive di default"],
        rows=[
            ["Edge & impianti",
             "Acquisizione telemetria, controllo locale, inferenza a bassa latenza su anomaly detection",
             "On-prem (centrali, parchi, substation), edge runtime K3s"],
            ["Private cloud",
             "Data hub on-prem, OSS ML platform per modelli sensibili, vault per segreti, identità unificata",
             "Data center primario EnergoGrid"],
            ["Public cloud (multi-cloud)",
             "Training elastico, batch scoring su grandi volumi, modelli generativi, analytics warehouse-scale",
             "Landing zone AWS, Azure, GCP secondo workload"],
            ["Consumer",
             "Cruscotti dispatcher, asset management, trading, reporting regolatorio",
             "Sistemi aziendali esistenti, integrati via API"],
        ],
        col_widths_cm=[3.5, 9.0, 4.0],
        first_col_bold=True,
    )

    # 5.3
    add_heading(doc, "Pattern di distribuzione dei workload",
                level=2, numbered_prefix="5.3")
    add_paragraph(
        doc,
        "Il framework decisionale che assegna ciascun workload al "
        "cluster corretto è volutamente semplice. Due sole domande, "
        "applicate in ordine, conducono a una raccomandazione iniziale; "
        "i guardrail trasversali assicurano che la scelta non comprometta "
        "i principi della sezione 5.1."
    )

    add_image(doc, "04_paas_iaas_oss_decision_tree.png",
              caption="Figura 5.3 - Framework decisionale per l'assegnazione "
                      "di un workload AI: quattro domande, cinque esiti (hybrid "
                      "burst, OSS on-prem, PaaS pieno, PaaS curato, IaaS+OSS).")

    add_paragraph(
        doc,
        "L'applicazione del framework ai tre casi d'uso EnergoGrid "
        "produce una distribuzione che evita la trappola del single-cloud "
        "e quella opposta del full multi-cloud replicato."
    )

    add_table(
        doc,
        headers=["Caso d'uso", "Componente principale", "Cluster di destinazione", "Motivazione"],
        rows=[
            ["Forecasting domanda",
             "Training",
             "GCP (BigQuery + Vertex AI)",
             "I dati di consumo e di mercato vivono nel warehouse, il training tabellare è nativo, latenza accettabile"],
            ["Forecasting domanda",
             "Serving online",
             "AWS SageMaker endpoint",
             "Coerenza con resto del programma ML, autoscaling, esposizione via API gateway esistente"],
            ["Manutenzione predittiva",
             "Training",
             "On-prem OSS",
             "Dati di telemetria sensibili, vincolo di sovranità, volumi gestibili da GPU pool interno"],
            ["Manutenzione predittiva",
             "Serving",
             "On-prem (batch) + edge (anomaly real-time)",
             "Batch scoring giornaliero notturno; inferenza edge a substation per anomaly su parametri elettrici"],
            ["Ottimizzazione bilanciamento",
             "Training",
             "On-prem OSS",
             "Logica core dell'operatore, vincolo di sovranità totale, no esposizione esterna"],
            ["Ottimizzazione bilanciamento",
             "Serving online",
             "On-prem OSS",
             "Latenza p95 sotto 500 ms verso dispatcher, integrazione diretta con SCADA"],
            ["Reportistica e documentazione",
             "LLM enterprise",
             "Azure OpenAI Service (regione EU)",
             "Garanzie contrattuali enterprise, integrazione con Office 365 e Teams"],
        ],
        col_widths_cm=[3.5, 3.5, 4.0, 5.5],
        first_col_bold=True,
    )

    add_paragraph(
        doc,
        "Tradotta in termini di lifecycle, l'assegnazione produce due "
        "lane di esecuzione che condividono lo stesso model registry "
        "(MLflow) come fonte unica di verità. La private lane esegue "
        "training su dati sensibili e mantiene endpoint a bassa latenza "
        "verso il dispatch; la public lane esegue training non sensibile "
        "su GPU spot e mantiene endpoint per workload elastici. La "
        "figura 5.4 esplicita come modelli e drift signal attraversano le "
        "lane senza che i dati sensibili lascino mai il perimetro privato."
    )

    add_image(doc, "05_ml_lifecycle_hybrid.png",
              caption="Figura 5.4 - Lifecycle ML ibrido. Private e public "
                      "lane condividono un unico model registry; la "
                      "promozione a produzione passa per approvazione "
                      "esplicita (promote-on-approval).",
              width_cm=16.0)

    # 5.4
    add_heading(doc, "Modello di responsabilità multi-cloud",
                level=2, numbered_prefix="5.4")
    add_paragraph(
        doc,
        "L'introduzione di tre cloud pubblici e dell'OSS on-prem moltiplica "
        "i punti di responsabilità. Una matrice esplicita di chi-fa-cosa "
        "è una precondizione operativa, non una formalità. La tabella "
        "sintetizza la responsabilità di alto livello sui domini "
        "principali. Il dettaglio operativo è rimandato ai runbook di "
        "ciascun workstream."
    )

    add_table(
        doc,
        headers=["Dominio", "Provider PaaS", "EnergoGrid IT", "EnergoGrid OT", "EnergoGrid ML"],
        rows=[
            ["Hardware fisico",       "Cloud",        "Data center", "OT plant",   "n/a"],
            ["Sistema operativo",     "Cloud (PaaS)", "On-prem",     "On-prem OT", "n/a"],
            ["Container runtime",     "Cloud",        "On-prem",     "Edge",       "n/a"],
            ["Dati grezzi",           "n/a",          "Pipeline",    "Sorgente",   "n/a"],
            ["Feature engineering",   "n/a",          "n/a",         "n/a",        "Owner"],
            ["Modelli e training",    "n/a",          "n/a",         "n/a",        "Owner"],
            ["Deployment endpoint",   "Runtime",      "Network",     "n/a",        "Codice"],
            ["Drift monitoring",      "Tooling",      "n/a",         "n/a",        "Owner"],
            ["Identità e accessi",   "IAM service",  "Owner",       "OT-IAM",     "Consumer"],
            ["Audit log",             "Servizio",     "Aggregazione","OT events",  "Eventi ML"],
            ["Compliance regolatoria","Certificazioni","Coordina",    "OT compliance","Modelli e dati"],
        ],
        col_widths_cm=[3.5, 3.0, 2.7, 3.0, 3.3],
        first_col_bold=True,
    )

    add_callout(
        doc,
        "implicazione",
        "La matrice mette in chiaro un punto facile da sottovalutare: "
        "ogni dominio ha un owner unico ma tre o quattro contributori. "
        "La governance va trattata come condizione di partenza: è ciò che "
        "impedisce alla responsabilità di evaporare nei punti di contatto "
        "fra cluster (sezione 7)."
    )

    _add_page_break(doc)


# === Section 6 - Business and cost impact ================================


def build_section_6(doc):
    add_heading(doc, "Impatto su processi e costi",
                level=1, numbered_prefix="6.")

    # 6.1
    add_heading(doc, "Impatto organizzativo", level=2, numbered_prefix="6.1")
    add_paragraph(
        doc,
        "L'architettura ibrida cambia il profilo di capacità del team. "
        "Un'organizzazione progettata per gestire infrastruttura "
        "esclusivamente on-prem deve evolvere in tre dimensioni: cloud "
        "engineering (FinOps, IaC, networking cross-cloud), MLOps "
        "(experiment management, model registry, drift), governance "
        "tecnica (IAM unificato, audit cross-cloud, compliance evidenze)."
    )

    add_table(
        doc,
        headers=["Capacità", "Stato iniziale", "Target a regime", "Modalità di build"],
        rows=[
            ["Cloud engineering / FinOps",
             "Limitato, sperimentale su AWS",
             "Team dedicato 4-6 FTE multi-cloud certificati",
             "Hiring esterno + percorso di certificazione interno"],
            ["MLOps / data engineering",
             "Modelli statistici legacy, no MLOps",
             "Squadra 6-8 FTE su MLflow, feature store, monitoring drift",
             "Hiring senior + formazione interna su giuniori"],
            ["Cybersecurity OT",
             "Solido sul perimetro tradizionale",
             "Esteso al perimetro cloud, IEC 62443 + zero trust",
             "Upskilling team esistente + consulenza specialistica"],
            ["Governance e compliance",
             "Funzione regolatoria classica",
             "Funzione tecnica con evidenze automatizzate audit-ready",
             "Estensione del team risk con un product manager dedicato"],
            ["Solution architecture",
             "Pratica concentrata su singoli sistemi",
             "Architetti cross-cloud con visione di programma",
             "Promozione interna + uno o due senior dall'esterno"],
        ],
        col_widths_cm=[3.5, 3.5, 4.5, 5.0],
        first_col_bold=True,
    )

    add_paragraph(
        doc,
        "Il piano di transizione è incrementale: la wave 1 della "
        "roadmap (sezione 8) è dimensionata per essere sostenibile dal "
        "team di partenza, con onboarding di rinforzi durante la wave 2 e "
        "stabilizzazione del modello operativo in wave 3. Forzare la "
        "trasformazione organizzativa prima dell'infrastruttura genera "
        "frustrazione e turnover; muoversi nella sequenza opposta fa "
        "imparare sul prodotto reale, non su slide."
    )

    # 6.2
    add_heading(doc, "TCO concettuale a 5 anni", level=2, numbered_prefix="6.2")
    add_paragraph(
        doc,
        "La proiezione TCO resta a livello concettuale. Non vengono "
        "stimati euro puntuali per voce: l'incertezza intrinseca su "
        "volumi, prezzi unitari cloud e ciclo di refresh hardware "
        "renderebbe una stima numerica dettagliata fuorviante. L'obiettivo "
        "è confrontare la forma del costo on-prem only contro l'hybrid "
        "multi-cloud sull'orizzonte 5 anni e identificare i driver "
        "principali."
    )

    add_image(doc, "06_tco_5y_projection.png",
              caption="Figura 6.1 - Proiezione concettuale del TCO a 5 anni. "
                      "Le unità sono proxy normalizzati: la grandezza "
                      "informativa è il rapporto fra le due curve, non "
                      "il valore assoluto.")

    add_paragraph(
        doc,
        "Tre osservazioni emergono dalla figura:"
    )
    add_bullets(doc, [
        "Anno 1: lo scenario hybrid è già più economico dello scenario "
        "on-prem only, perché evita il refresh hardware massiccio per "
        "GPU di training. Il differenziale (~30%) è il maggior beneficio "
        "del periodo di startup.",

        "Anni 2-4: la curva on-prem ha un creep operativo lento ma "
        "costante (ricambio licenze, manutenzione, consumo energetico); "
        "la curva hybrid è praticamente piatta, perché il costo "
        "variabile cloud assorbe l'incremento di volume mentre la base "
        "on-prem resta dimensionata sul perimetro sensibile.",

        "Anno 5: il differenziale cumulato si attesta intorno al 20%. "
        "Non è il numero da raccontare al CFO come unica metrica: il "
        "valore vero è la differente flessibilità. L'hybrid permette "
        "di assorbire un raddoppio del volume di forecasting senza "
        "acquisti CAPEX; l'on-prem only richiederebbe un altro refresh.",
    ])

    add_paragraph(
        doc,
        "L'hybrid porta driver di costo nascosti noti dalla letteratura "
        "cloud: egress cross-cloud, tooling di osservabilità unificato, "
        "operations multi-cloud, perdita di volume discount e GPU idle in "
        "training. Sono coperti esplicitamente nella mappa dei rischi "
        "(sezione 8.3), che ne riporta mitigazione e owner."
    )

    _add_page_break(doc)


# === Section 7 - Security, compliance, data governance ===================


def build_section_7(doc):
    add_heading(doc, "Sicurezza, compliance e data governance",
                level=1, numbered_prefix="7.")

    add_paragraph(
        doc,
        "Per un operatore di rete energetica la sicurezza non è un "
        "modulo aggiuntivo: è il contesto in cui ogni altra scelta vive. "
        "La sezione è organizzata per scendere progressivamente dal "
        "quadro normativo applicabile, ai controlli per dominio, alla "
        "data governance, fino alla checklist di alto livello che "
        "EnergoGrid può usare come baseline di valutazione interna."
    )

    # 7.1
    add_heading(doc, "Quadro normativo applicabile", level=2, numbered_prefix="7.1")
    add_paragraph(
        doc,
        "Le normative rilevanti per EnergoGrid che impattano "
        "l'architettura della piattaforma AI sono cinque, ciascuna con "
        "implicazioni diverse sul disegno e sui controlli da prevedere."
    )

    add_table(
        doc,
        headers=["Norma", "Ambito", "Implicazione architetturale"],
        rows=[
            ["GDPR (Reg. UE 2016/679)",
             "Dati personali di clienti e dipendenti",
             "Data residency EU, accesso controllato, diritto alla cancellazione, registro trattamenti"],
            ["NIS2 (Dir. UE 2022/2555)",
             "Operatori essenziali, sicurezza cyber e resilienza",
             "Notifica incidenti entro 24h, gestione vulnerabilità, SBOM, piano DR testato"],
            ["EU AI Act (Reg. UE 2024/1689)",
             "Sistemi AI ad alto rischio, gestione e trasparenza",
             "Sistemi di gestione del rischio, conservazione log, supervisione umana sui sistemi critici"],
            ["ISO/IEC 27001 + 27019",
             "ISMS generico + specifico settore energia",
             "Gestione asset, controlli su accesso, fornitori, business continuity"],
            ["IEC 62443",
             "Cybersecurity per industrial automation",
             "Segmentazione zone/conduit, security level per zona, hardening componenti OT"],
        ],
        col_widths_cm=[3.5, 4.0, 8.5],
        first_col_bold=True,
    )

    add_paragraph(
        doc,
        "EnergoGrid rientra come operatore essenziale sotto NIS2 (settore "
        "energia, sottosettore elettricità); l'attivazione AI più "
        "sensibile (ottimizzazione dispacciamento) ricade nel perimetro "
        "AI Act come sistema potenzialmente ad alto rischio (gestione di "
        "infrastruttura critica). Questi due punti vincolano l'architettura "
        "più di GDPR e ISO, che sono prerequisiti già coperti dai "
        "processi aziendali esistenti."
    )

    # 7.2
    add_heading(doc, "Controlli per dominio", level=2, numbered_prefix="7.2")
    add_paragraph(
        doc,
        "I controlli sono organizzati per dominio funzionale: identità, "
        "dato, modello, rete, edge, audit. La figura 7.1 sintetizza la "
        "mappa, evidenziando i principi al centro (least privilege, zero "
        "trust, data sovereignty) e l'inquadramento normativo agli angoli."
    )

    add_image(doc, "10_governance_domain_map.png",
              caption="Figura 7.1 - Mappa dei domini di governance "
                      "multi-cloud. Fascia dei principi in alto, sei domini "
                      "nella griglia centrale, norme applicabili nella fascia "
                      "inferiore.")

    add_table(
        doc,
        headers=["Dominio", "Controlli chiave"],
        rows=[
            ["Identità e accessi",
             "IAM federato Entra ID + provider cloud, MFA obbligatorio, role-based access con review trimestrale, separazione IT vs OT a livello di identità"],
            ["Dato",
             "Classificazione esplicita (pubblico, interno, confidenziale, riservato OT), cifratura at-rest e in-transit con chiavi gestite via HashiCorp Vault, data lineage end-to-end"],
            ["Modello",
             "Model registry centralizzato (MLflow), approvazione esplicita per promozione a produzione, model card per ogni versione, evaluator di bias e fairness dove applicabile"],
            ["Rete",
             "Segmentazione zone IEC 62443, private link verso ogni cloud, no esposizione pubblica di endpoint operativi, ispezione del traffico cross-cloud"],
            ["Edge",
             "Device identity con certificati hardware, OTA firmati e verificati, hardening del runtime K3s, monitoring di tampering fisico"],
            ["Audit",
             "Log centralizzato (siem) con immutabilità WORM, retention 5+ anni, alerting su pattern anomali, evidenze esportabili in formato audit-ready"],
        ],
        col_widths_cm=[3.0, 13.5],
        first_col_bold=True,
    )

    # 7.3
    add_heading(doc, "Data governance e lineage", level=2, numbered_prefix="7.3")
    add_paragraph(
        doc,
        "La data governance è la disciplina che evita che la "
        "piattaforma diventi un buco nero in cui si vede entrare dato e "
        "uscire previsioni, senza tracciabilità di cosa è successo nel "
        "mezzo. Per EnergoGrid quattro pratiche sono non negoziabili."
    )
    add_bullets(doc, [
        "Catalogo dati centrale con classificazione esplicita e owner "
        "documentato per ciascun dataset, sia all'origine sia nelle "
        "trasformazioni intermedie.",

        "Data lineage automatico, tracciato dalle pipeline di ingestione "
        "fino alle feature usate dai modelli; ogni previsione è "
        "ricostruibile a ritroso fino ai dati grezzi che l'hanno prodotta.",

        "Feature store come unica fonte di verità: nessuna feature viene "
        "calcolata in app code; ogni feature ha definizione, owner, "
        "criteri di freshness e signature.",

        "Conservazione retention-aware: ogni dataset ha politica di "
        "retention esplicita (tipicamente 5 anni per i dati operativi "
        "che entrano in modelli ad uso regolato, durata di vita asset per "
        "la telemetria critica), automatizzata via lifecycle policy.",
    ])

    # 7.4
    add_heading(doc, "Checklist di alto livello", level=2, numbered_prefix="7.4")
    add_paragraph(
        doc,
        "La checklist sintetizza i requisiti minimi di compliance e "
        "sicurezza che EnergoGrid deve garantire prima del passaggio in "
        "produzione di ogni workload AI. Non sostituisce gli audit "
        "regolatori; serve come gate operativo interno e come traccia per "
        "le evidenze."
    )

    add_table(
        doc,
        headers=["Area", "Requisito minimo", "Riferimento"],
        rows=[
            ["Residency",     "Dato di telemetria e PII in regione EU, contratto sub-processor verificato", "GDPR, NIS2"],
            ["Identità",     "MFA + accesso condizionale per ogni utente, separazione IT-OT", "ISO 27001, IEC 62443"],
            ["Cifratura",     "At-rest AES-256 con chiavi in Vault, in-transit TLS 1.3, rotazione chiavi automatica", "ISO 27001, NIS2"],
            ["Audit",         "Log immutabili 5+ anni, evidenza accessi privilegiati, alert su anomalie", "NIS2, EU AI Act"],
            ["Modello",       "Model card, valutazione bias dove applicabile, approvazione esplicita per prod", "EU AI Act"],
            ["DR",            "Piano DR documentato, RTO < 4h, RPO < 1h, test annuale", "NIS2, ISO 27001"],
            ["Vulnerability", "Patch SLA chiaro, SBOM per ogni componente, scanning automatico CI/CD", "NIS2"],
            ["Incident",      "Notifica regolatore entro 24h, runbook documentati, exercise annuale", "NIS2"],
            ["Fornitori",     "Due diligence sub-processor, valutazione sicurezza periodica, exit strategy", "ISO 27001, NIS2"],
            ["Supervisione",  "Human-in-the-loop per decisioni di dispacciamento, override sempre disponibile", "EU AI Act"],
        ],
        col_widths_cm=[2.5, 11.0, 3.0],
        first_col_bold=True,
    )

    add_callout(
        doc,
        "nota",
        "Le voci della checklist sono di alto livello: ciascuna si traduce "
        "in decine di controlli specifici nei framework di riferimento. "
        "L'obiettivo qui non è fare la mappatura dettagliata, ma "
        "garantire che ogni area sia presidiata fin dalla fase di disegno "
        "e non lasciata come debito da risolvere dopo il go-live."
    )

    _add_page_break(doc)


# === Section 8 - Migration roadmap =======================================


def build_section_8(doc):
    add_heading(doc, "Piano di migrazione e roadmap",
                level=1, numbered_prefix="8.")

    add_paragraph(
        doc,
        "La roadmap copre 36 mesi ed è organizzata in tre wave "
        "sequenziali ma sovrapposte: ogni wave entra in esecuzione mentre "
        "la precedente passa a stabilizzazione. La scelta di un orizzonte "
        "triennale è deliberata: la consegna privilegia profondità "
        "strategica e visibilità regolatoria piuttosto che il rilascio "
        "rapido di una capacità isolata. Una roadmap a 12 mesi sarebbe "
        "stata adeguata per un MVP, ma non per un programma che ridisegna "
        "il modello operativo AI di un operatore essenziale."
    )

    # 8.1
    add_heading(doc, "Strategia incrementale a tre wave",
                level=2, numbered_prefix="8.1")
    add_image(doc, "07_migration_roadmap_waves.png",
              caption="Figura 8.1 - Roadmap di migrazione a tre wave su 36 mesi, "
                      "con tre workstream paralleli (infra & governance, ML "
                      "platform & lifecycle, use case delivery), sei milestone "
                      "su timeline e tre gate strategici.",
              width_cm=16.0)

    add_table(
        doc,
        headers=["Wave", "Mesi", "Obiettivo strategico", "Output principali"],
        rows=[
            ["Wave 1 · Foundations",
             "M0-M9",
             "Costruire la fondazione: landing zone multi-cloud, governance baseline, data hub on-prem, primo pilota in produzione",
             "Landing zone su 1 cloud + on-prem, IAM federato, OSS ML platform v1, pilota forecasting su nodo rappresentativo"],
            ["Wave 2 · Hybrid lifecycle",
             "M6-M21",
             "Industrializzare il ciclo di vita ML: feature store, registry, drift monitoring, automazione retraining",
             "MLflow operativo, feature store condiviso, pred. maintenance in produzione, secondo cloud onboarded"],
            ["Wave 3 · Scale & optimisation",
             "M18-M36",
             "Scalare il programma, ottimizzare costo e affidabilità, estendere edge inference",
             "Dispatch optimisation in prod, edge runtime a substation pilota, multi-region per DR, ottimizzazione TCO"],
        ],
        col_widths_cm=[3.5, 1.8, 5.0, 6.2],
        first_col_bold=True,
    )

    add_paragraph(
        doc,
        "La sovrapposizione fra wave è voluta: la wave successiva ha "
        "bisogno di alcune fondamenta dalla precedente, e sospendere lo "
        "sviluppo per attendere la chiusura formale costerebbe calendario "
        "senza ridurre il rischio. La gestione di "
        "questa sovrapposizione è la responsabilità del program manager "
        "del capstone, supportata dai gate strategici dove le decisioni "
        "di proseguire o riallineare vengono prese formalmente."
    )

    # 8.2
    add_heading(doc, "Milestone e gate strategici",
                level=2, numbered_prefix="8.2")

    add_table(
        doc,
        headers=["ID", "Mese", "Milestone", "Cosa significa raggiungerla"],
        rows=[
            ["M1", "M3",  "Landing zone operativa",
             "Primo cloud configurato con guardrail di sicurezza, network, IAM federato; sub-processor approvati"],
            ["M2", "M9",  "Pilota forecasting live",
             "Modello forecasting in produzione su nodo pilota, output integrato al dispatcher, drift monitoring attivo"],
            ["M3", "M15", "Feature store condiviso",
             "Feature store operativo, almeno 30 feature documentate, condiviso tra forecasting e manutenzione predittiva"],
            ["M4", "M21", "Pred. maint. in produzione",
             "Modello manutenzione predittiva su flotta inverter completa, integrato con asset management, retraining automatico"],
            ["M5", "M27", "Dispatch optimisation",
             "Algoritmo di ottimizzazione bilanciamento in produzione con supervisione umana, gate AI Act soddisfatto"],
            ["M6", "M36", "Regime operativo",
             "Tre casi d'uso in regime, OPEX stabilizzato, KPI di programma raggiunti, modello operativo replicabile"],
        ],
        col_widths_cm=[1.0, 1.5, 4.0, 10.0],
        first_col_bold=True,
    )

    add_paragraph(
        doc,
        "I tre gate strategici sono punti di decisione formale, non "
        "checkpoint informativi. La decisione richiede sponsorship "
        "esplicito del comitato di programma e ha tre esiti possibili: "
        "procedere, riallineare scope o budget, fermare il programma."
    )

    add_table(
        doc,
        headers=["Gate", "Mese", "Decisione", "Criteri di valutazione"],
        rows=[
            ["G1 · Go/No-Go pilot",
             "M3",
             "Procedere oltre la wave 1?",
             "Landing zone sicura, primo modello in training in cloud, NO trovati sulla compliance baseline"],
            ["G2 · Approval scaling",
             "M15",
             "Sbloccare la wave 2?",
             "Pilota forecasting in produzione stabile da almeno 60 giorni, KPI MAPE entro target, budget M0-M15 nei limiti"],
            ["G3 · Vendor review",
             "M27",
             "Confermare o ribilanciare il mix cloud?",
             "Analisi TCO reale vs forecast, valutazione concentrazione spesa, opportunità di consolidamento, contratto pluriennale"],
        ],
        col_widths_cm=[3.5, 1.5, 4.0, 7.5],
        first_col_bold=True,
    )

    # 8.3
    add_heading(doc, "Rischi principali e mitigazioni",
                level=2, numbered_prefix="8.3")
    add_paragraph(
        doc,
        "I rischi principali sono stati posizionati su una mappa "
        "probabilità x impatto. La zona rossa (probabilità o impatto "
        "alti) raccoglie i rischi che hanno una mitigazione attiva nella "
        "wave 1 della roadmap, perché lasciarli aperti oltre il primo "
        "gate avrebbe ripercussioni a cascata."
    )

    add_image(doc, "08_risk_heatmap.png",
              caption="Figura 8.2 - Mappa probabilità x impatto dei rischi "
                      "principali. La zona rossa richiede mitigazione "
                      "esplicita nella roadmap.")

    add_table(
        doc,
        headers=["ID", "Rischio", "Mitigazione raccomandata", "Owner"],
        rows=[
            ["R1", "Cross-cloud egress costs",
             "Co-locazione regionale, batch instead of streaming dove possibile, monitoraggio dedicato",
             "FinOps + Cloud Engineering"],
            ["R2", "Skill gap MLOps",
             "Piano hiring senior in wave 1, percorso certificazione interna, partnership formazione con vendor",
             "HR + ML Engineering"],
            ["R3", "Vendor lock-in PaaS",
             "Astrazione interna obbligatoria, uso preferenziale di componenti OSS, exit strategy documentata",
             "Solution Architecture"],
            ["R4", "Drift modelli non rilevato",
             "Drift monitoring automatico da wave 1, alert su SLO, runbook di retraining",
             "ML Engineering"],
            ["R5", "Incidente cyber su SCADA",
             "Segmentazione IT-OT rigorosa, IEC 62443 baseline, exercise periodici, SOC dedicato OT",
             "Cybersecurity OT"],
            ["R6", "Lentezza interop con sistemi legacy",
             "POC di integrazione in wave 1, abstraction layer dedicato, fallback a integrazione asincrona",
             "Solution Architecture + IT legacy"],
            ["R7", "Audit normativo NIS2",
             "Evidenze automatizzate fin dalla wave 1, mock audit semestrale, relazione con autorità competente",
             "Governance & Compliance"],
            ["R8", "Adozione utenti operativi insufficiente",
             "Coinvolgimento dispatcher dalla fase di design, change management strutturato, formazione continua",
             "Operations + Change Management"],
            ["R9", "Costo unbudgeted GPU cloud",
             "Quota cloud per progetto, spot/preemptible per training non urgente, alert anomali su billing",
             "FinOps"],
            ["R10", "Failover cross-region che non funziona",
             "DR drill annuale (chaos engineering), runbook automatizzati, validazione RTO/RPO via test",
             "Cloud Engineering + Operations"],
        ],
        col_widths_cm=[1.0, 4.0, 7.0, 4.5],
        first_col_bold=True,
    )

    _add_page_break(doc)


# === Section 9 - KPI and measurement =====================================


def build_section_9(doc):
    add_heading(doc, "KPI di progetto e misurazione",
                level=1, numbered_prefix="9.")

    add_paragraph(
        doc,
        "I KPI sono organizzati in albero gerarchico, con il valore di "
        "business alla radice e le metriche operative misurabili come "
        "foglie. La gerarchia evita due errori frequenti: misurare solo "
        "metriche di output (precision, recall) senza collegarle al "
        "business, e raccontare risultati di business non riconducibili a "
        "metriche misurabili. Il sistema di misurazione è parte "
        "integrante della piattaforma, non una funzione di reportistica "
        "successiva."
    )

    add_image(doc, "09_kpi_tree.png",
              caption="Figura 9.1 - Albero dei KPI di progetto. Valore di "
                      "business in radice, tre KPI macro intermedi, sei "
                      "metriche operative alle foglie.")

    add_heading(doc, "KPI di business", level=2, numbered_prefix="9.1")
    add_table(
        doc,
        headers=["KPI", "Definizione", "Target a 24 mesi", "Baseline"],
        rows=[
            ["MAPE forecasting domanda",
             "Mean Absolute Percentage Error su forecast T+24h aggregato per macro-nodo",
             "<= 4%",
             "~7% (modelli statistici legacy)"],
            ["MTBF asset rotanti",
             "Mean Time Between Failures su inverter e trasformatori monitorati",
             "+15% rispetto a baseline",
             "Misurazione attuale da asset management"],
            ["Costo bilanciamento (EUR/MWh)",
             "Costo medio per MWh di sbilanciamento attribuibile a previsione/ottimizzazione",
             "-8% rispetto a baseline",
             "Costo storico ultimi 24 mesi"],
            ["Riduzione downtime non pianificato",
             "Ore di downtime non pianificato su asset monitorati, media annuale",
             "-20% in 24 mesi",
             "Registro incidenti operativi"],
            ["Penetrazione rinnovabili gestibile",
             "Quota di rinnovabili che la rete può assorbire senza intervento manuale di stabilizzazione",
             "+10 punti percentuali",
             "Misurazione attuale dispatcher"],
        ],
        col_widths_cm=[4.0, 7.0, 3.0, 2.5],
        first_col_bold=True,
    )

    add_heading(doc, "KPI tecnici e operativi (MLOps)", level=2,
                numbered_prefix="9.2")
    add_table(
        doc,
        headers=["KPI", "Definizione", "Target"],
        rows=[
            ["Time-to-deploy modello",
             "Tempo dalla decisione di promozione alla disponibilità in produzione",
             "<= 2 settimane"],
            ["Frequenza retraining",
             "Cadenza media di retraining automatico per ciascun modello in produzione",
             "settimanale o triggered da drift"],
            ["Tasso di rollback",
             "Percentuale di modelli che vengono rolled-back entro 30 giorni dal deploy",
             "<= 5%"],
            ["Drift detection latency",
             "Tempo medio fra l'insorgenza di drift e la generazione dell'alert",
             "<= 24 ore"],
            ["Endpoint availability",
             "Disponibilità degli endpoint di serving online",
             ">= 99.9%"],
            ["Inferenza p95 (online)",
             "Latenza percentile 95 dei serving endpoint online",
             "<= 500 ms"],
            ["Inferenza p95 (edge)",
             "Latenza percentile 95 dell'anomaly detection a substation",
             "<= 50 ms"],
        ],
        col_widths_cm=[4.5, 8.5, 3.5],
        first_col_bold=True,
    )

    add_heading(doc, "KPI di governance e compliance", level=2,
                numbered_prefix="9.3")
    add_table(
        doc,
        headers=["KPI", "Definizione", "Target"],
        rows=[
            ["Gap NIS2 aperti (critici)", "Numero di gap critici di compliance NIS2 non ancora chiusi", "0 dopo M9, mantenuto"],
            ["Audit findings (critici)",  "Findings di audit interno o esterno classificati critici",     "0 in regime"],
            ["Coverage audit log",         "Percentuale di azioni privilegiate coperte da log centralizzato", ">= 100% (a regime)"],
            ["Incidenti cyber su perimetro OT", "Numero di incidenti cyber con impatto su perimetro OT", "0 in 12 mesi continuativi"],
            ["Time-to-notification NIS2",  "Tempo dall'incidente alla notifica al regolatore",              "<= 24 ore"],
            ["Model card coverage",        "Percentuale di modelli in produzione dotati di model card aggiornata", ">= 100%"],
            ["Bias review coverage",       "Percentuale di modelli applicabili (impatto su persone/decisioni operative) con review bias annuale", ">= 100%"],
        ],
        col_widths_cm=[4.5, 8.5, 3.5],
        first_col_bold=True,
    )

    add_heading(doc, "Sistema di misurazione", level=2,
                numbered_prefix="9.4")
    add_paragraph(
        doc,
        "I KPI sono raccolti automaticamente dalla piattaforma: la "
        "reportistica manuale è bandita perché non scala e perché "
        "introduce ritardo nella reazione. Tre meccanismi coprono la "
        "raccolta."
    )
    add_bullets(doc, [
        "Per i KPI di business: pipeline batch giornaliera che aggrega le "
        "previsioni e i risultati osservati nel data warehouse, calcola "
        "il delta e lo espone in dashboard Grafana / Power BI.",

        "Per i KPI tecnici e operativi: Prometheus + MLflow estraggono "
        "metriche dai servizi e dai run di training; le serie temporali "
        "alimentano dashboard di servizio.",

        "Per i KPI di governance: combinazione di evidenze dal SIEM (per "
        "audit log e incidenti), dal model registry (per coverage model "
        "card e bias review) e dal sistema di ticketing (per findings di "
        "audit).",
    ])

    add_callout(
        doc,
        "principio",
        "Un KPI che richiede intervento manuale ogni mese per essere "
        "aggiornato non è un KPI, è un report. Tutti i target esposti "
        "in questa sezione sono pensati per essere calcolati e visibili "
        "in tempo quasi-reale dal sistema, non riassunti a posteriori."
    )

    _add_page_break(doc)


# === Section 10 - Conclusions and recommendations ========================


def build_section_10(doc):
    add_heading(doc, "Conclusioni e raccomandazioni strategiche",
                level=1, numbered_prefix="10.")

    add_paragraph(
        doc,
        "Il documento ha argomentato che per EnergoGrid la risposta "
        "corretta al problema della piattaforma AI non è un cloud "
        "preferito né un'architettura single-vendor portata all'estremo. "
        "E' una architettura ibrida costruita per assegnare ogni workload "
        "al cluster di soluzioni dove la sua intersezione di vincoli (di "
        "sovranità, latenza, volume, costo, integrazione) rende l'opzione "
        "dominante. La complessità che ne consegue è un costo reale, "
        "non un vezzo, e va gestito esplicitamente attraverso governance, "
        "automazione e disciplina di lock-in. In cambio, la flessibilità "
        "operativa, la resilienza ai cambi di scenario regolatorio e la "
        "riduzione del rischio tecnologico giustificano la scelta."
    )

    add_heading(doc, "Sintesi delle scelte strategiche", level=2,
                numbered_prefix="10.1")
    add_table(
        doc,
        headers=["Dimensione strategica", "Scelta raccomandata", "Motivazione sintetica"],
        rows=[
            ["Architettura complessiva",
             "Hybrid multi-cloud disciplinato",
             "Massimizza adempimento dei requisiti regolatori e flessibilità operativa, evita il full multi-cloud replicato"],
            ["Spina dorsale ML",
             "AWS SageMaker per workload elastici, OSS on-prem per sensibili",
             "Maturità della piattaforma, ecosistema, integrazione IoT industriale"],
            ["Analytics su grandi volumi",
             "GCP BigQuery + BigQuery ML",
             "Unico provider con warehouse serverless petabyte-scale e training tabellare nativo"],
            ["LLM enterprise",
             "Azure OpenAI Service in regione EU",
             "Garanzie contrattuali enterprise, integrazione Office 365, compliance certificata"],
            ["Componenti sensibili",
             "OSS on-prem (Kubernetes + MLflow + FastAPI)",
             "Sovranità dato, latenza, controllo del codice per audit, riduzione lock-in"],
            ["Identità e accessi",
             "Federazione cross-cloud su Entra ID",
             "Single source of truth, separazione IT-OT, MFA obbligatorio"],
            ["Osservabilità",
             "Prometheus + Grafana come base, integrazione provider-specific complementare",
             "Layer OSS portabile, copre cloud e on-prem in modo omogeneo"],
            ["Lifecycle ML",
             "MLflow come registry unico, cross-cluster",
             "Standard de facto, indipendente da provider, abilita la portabilità"],
        ],
        col_widths_cm=[4.0, 5.5, 7.0],
        first_col_bold=True,
    )

    add_heading(doc, "Priorità raccomandate", level=2,
                numbered_prefix="10.2")
    add_paragraph(
        doc,
        "Le tre priorità di intervento per i prossimi 12 mesi, in "
        "ordine di criticità rispetto al successo del programma."
    )
    add_bullets(doc, [
        "Costruire la fondazione di governance prima dell'infrastruttura. "
        "IAM federato, model registry, classificazione dato e quadro "
        "normativo vanno definiti e operativi entro M6, perché "
        "costruire ML su una governance traballante moltiplica il debito "
        "tecnico futuro.",

        "Investire nel team prima della tecnologia. Una piattaforma "
        "ibrida che il team non sa operare diventa un rischio strategico "
        "a tutti gli effetti. Hiring senior in wave 1, formazione strutturata, "
        "partnership con system integrator solo per accelerare, non per "
        "sostituire la capacità interna.",

        "Lanciare un solo pilota end-to-end in wave 1, non tre paralleli. "
        "Il pilota forecasting (volumi prevedibili, dato meno sensibile) è "
        "il candidato migliore perché valida la fondazione tecnica e "
        "organizzativa con rischio contenuto, e prepara il terreno per i "
        "due casi d'uso più impegnativi.",
    ])

    add_heading(doc, "Prossimi passi strategici", level=2,
                numbered_prefix="10.3")
    add_paragraph(
        doc,
        "Sul fronte decisionale, due valutazioni vanno chiuse nei "
        "prossimi 90 giorni perché bloccano l'avvio operativo: "
        "selezione del cloud primario fra AWS e Azure (la consegna "
        "raccomanda AWS sulla base della maturità ML, ma una valutazione "
        "specifica dei contratti enterprise esistenti può invertire la "
        "raccomandazione), e definizione formale del modello di "
        "ownership cross-funzionale fra IT, OT, ML e Compliance."
    )
    add_paragraph(
        doc,
        "Sul fronte operativo, tre azioni preparano la wave 1: kickoff "
        "del program management con sponsorship esecutiva esplicita, "
        "ingaggio di un partner di sistema per accelerare la landing zone "
        "e il primo cloud, avvio del piano hiring per le quattro figure "
        "senior identificate nella sezione 6.1."
    )
    add_paragraph(
        doc,
        "Sul fronte istituzionale, l'apertura formale del dialogo con il "
        "regolatore (ARERA per il mercato italiano, ACER a livello "
        "europeo) e con l'autorità competente NIS2 va anticipata. "
        "Comunicare il programma in fase di disegno è meno costoso che "
        "subire un'ispezione su un sistema già in produzione."
    )

    add_callout(
        doc,
        "chiusura",
        "L'AI nel settore energetico non è una corsa tecnologica: è una "
        "trasformazione operativa che ha l'AI come abilitatore. La "
        "piattaforma raccomandata in questo documento non è "
        "best-in-class su nessun singolo asse; è la combinazione che "
        "meglio bilancia i vincoli che EnergoGrid affronta nei prossimi "
        "cinque anni. La revisione periodica di questa scelta, a ogni "
        "gate, è la disciplina che la mantiene rilevante."
    )

    _add_page_break(doc)


# === Annex A - Glossary ==================================================


def build_annex_a(doc):
    add_heading(doc, "Allegato A · Glossario", level=1)

    glossary = [
        ("AIaaS",       "AI-as-a-Service. Modello di consumo di capacità AI come servizio cloud, fatturato a consumo."),
        ("ARERA",       "Autorità di Regolazione per Energia Reti e Ambiente, regolatore italiano del mercato elettrico e gas."),
        ("AutoML",      "Categoria di servizi che automatizzano la selezione di algoritmo e iperparametri su dati strutturati."),
        ("Azure ML",    "Azure Machine Learning. Piattaforma Microsoft per il ciclo di vita ML."),
        ("Azure OpenAI", "Servizio Azure che espone i modelli OpenAI sotto contratto enterprise con garanzie di non-training e residency."),
        ("BigQuery",    "Data warehouse serverless e columnar di Google Cloud, fatturato per byte scansionati."),
        ("CapEx / OpEx", "Capital Expenditure (investimento ammortizzato) vs Operating Expenditure (costo operativo corrente)."),
        ("CT",          "Continuous Training. Pratica MLOps di retraining automatico su trigger di drift o cadenza."),
        ("Data lineage", "Tracciabilità del percorso di un dato dalla sorgente fino al consumo, attraverso le trasformazioni intermedie."),
        ("Data sovereignty", "Vincolo regolatorio che impone la residenza del dato in una giurisdizione specifica."),
        ("Drift",       "Cambiamento della distribuzione dei dati di input (data drift) o della relazione fra input e target (concept drift)."),
        ("EU AI Act",   "Regolamento UE 2024/1689 sulla intelligenza artificiale, applicabile per fasi dal 2025."),
        ("Feature store", "Repository centralizzato di feature ML, con definizioni condivise fra training e serving."),
        ("FastAPI",     "Framework Python per la realizzazione di API REST asincrone con tipizzazione forte e OpenAPI nativo."),
        ("FinOps",      "Disciplina di gestione finanziaria del cloud (governance del consumo, ottimizzazione, accountability)."),
        ("GDPR",        "Regolamento UE 2016/679 sulla protezione dei dati personali."),
        ("HIL",         "Human-in-the-loop. Supervisione umana sulle decisioni di un sistema automatico."),
        ("Hybrid",      "Architettura che combina modelli di servizio diversi (PaaS, IaaS, OSS) o cloud diversi."),
        ("IaaS",        "Infrastructure-as-a-Service. Modello cloud in cui il provider fornisce risorse infrastrutturali (compute, storage, rete)."),
        ("IEC 62443",   "Standard di cybersecurity per industrial automation and control systems."),
        ("ISO 27001",   "Standard internazionale per sistemi di gestione della sicurezza delle informazioni (ISMS)."),
        ("ISO 27019",   "Estensione settoriale di ISO 27001 per il settore dell'energia."),
        ("Lock-in",     "Costo e attrito di sostituire un fornitore con un altro; più alto in PaaS proprietari, più basso in OSS."),
        ("MLflow",      "Piattaforma open source per experiment tracking, model registry e gestione del lifecycle ML."),
        ("MLOps",       "Disciplina che applica i principi DevOps al ciclo di vita ML (CI, CD, CT, monitoring)."),
        ("MTBF",        "Mean Time Between Failures. Indicatore di affidabilità degli asset."),
        ("MAPE",        "Mean Absolute Percentage Error. Metrica di errore relativo per modelli di forecasting."),
        ("NIS2",        "Direttiva UE 2022/2555 sulla sicurezza delle reti e dei sistemi informativi, recepita in Italia nel 2024."),
        ("OPC UA",      "Open Platform Communications Unified Architecture. Protocollo standard di comunicazione industriale."),
        ("OSS",         "Open Source Software."),
        ("OT",          "Operational Technology. Sistemi di automazione industriale e controllo di processo, distinti dall'IT tradizionale."),
        ("PaaS",        "Platform-as-a-Service. Modello cloud in cui il provider gestisce la piattaforma applicativa, il consumer porta solo il codice e i dati."),
        ("PII",         "Personally Identifiable Information."),
        ("Prosumer",    "Utente che è contemporaneamente consumatore e produttore di energia (tipicamente fotovoltaico domestico)."),
        ("RAG",         "Retrieval-Augmented Generation. Pattern di LLM che attinge a una base di conoscenza esterna durante l'inferenza."),
        ("RACI",        "Modello di riferimento per le responsabilità (Responsible, Accountable, Consulted, Informed); la matrice della sezione 5.4 ne adotta una variante semplificata."),
        ("RPO / RTO",   "Recovery Point Objective / Recovery Time Objective. Metriche di disaster recovery."),
        ("SageMaker",   "Piattaforma ML end-to-end di AWS."),
        ("SCADA",       "Supervisory Control And Data Acquisition. Sistema di supervisione e controllo dei processi industriali."),
        ("Spot / preemptible", "Istanze cloud a costo ridotto ma con possibilità di revoca; adatte a workload tolerant a interruzione."),
        ("Substation",  "Sottostazione elettrica. Nodo di trasformazione e smistamento sulla rete."),
        ("TCO",         "Total Cost of Ownership. Costo complessivo di possesso di un sistema su un orizzonte definito."),
        ("Vertex AI",   "Piattaforma ML end-to-end di Google Cloud."),
        ("Zero trust",  "Modello di sicurezza che non assume fiducia implicita interna alla rete; ogni accesso è verificato esplicitamente."),
    ]

    add_table(
        doc,
        headers=["Termine", "Definizione"],
        rows=glossary,
        col_widths_cm=[3.0, 13.5],
        first_col_bold=True,
    )

    _add_page_break(doc)


# === Annex B - Technical references ======================================


def build_annex_b(doc):
    add_heading(doc, "Allegato B · Riferimenti tecnici", level=1)

    add_paragraph(
        doc,
        "Riferimenti consultati durante la stesura del documento. Le note "
        "del modulo 05 del Master sono la fonte primaria, integrate dalla "
        "documentazione ufficiale dei provider e da pubblicazioni della "
        "comunità MLOps."
    )

    add_heading(doc, "Materiale didattico del modulo", level=2)
    add_bullets(doc, [
        "Modulo 05 - Note 01: AI-as-a-Service and Cloud Architecture Fundamentals.",
        "Modulo 05 - Note 02: AWS AI and ML Stack.",
        "Modulo 05 - Note 03: Azure AI Ecosystem.",
        "Modulo 05 - Note 04: Google Cloud Vertex AI Data-First.",
        "Modulo 05 - Note 05: IaaS, Open Source AI, and On-Premises Deployment.",
        "Modulo 05 - Note 06: PaaS vs IaaS vs OSS Decision Framework.",
        "Modulo 05 - Note 07: Hybrid and Multi-Cloud Architecture Patterns.",
        "Modulo 04 - Note 01: Identifying AI Problems and Feasibility (TCO 3 anni).",
        "Modulo 04 - Note 02: KPIs, Lifecycle and Drift.",
        "Modulo 03 - Note 07: Agent Deployment (pattern di serving in container).",
    ])

    add_heading(doc, "Documentazione vendor", level=2)
    add_bullets(doc, [
        "AWS Well-Architected Framework, in particolare Machine Learning Lens.",
        "Microsoft Cloud Adoption Framework for Azure, sezione AI Strategy.",
        "Google Cloud Architecture Framework, in particolare AI and ML perspective.",
        "Kubernetes (kubernetes.io) e K3s (k3s.io) per edge orchestration.",
        "MLflow documentation (mlflow.org) per tracking e model registry.",
        "Hugging Face Hub (huggingface.co) per supply chain dei modelli pre-addestrati.",
        "FastAPI documentation (fastapi.tiangolo.com) per il pattern di serving Python.",
    ])

    add_heading(doc, "Pratiche di settore", level=2)
    add_bullets(doc, [
        "CNCF Cloud Native Glossary, per terminologia condivisa cloud-native.",
        "ENISA Cloud Security Guide per il settore energia.",
        "NIST AI Risk Management Framework (NIST AI 100-1), riferimento volontario per la gestione del rischio AI.",
        "Linee guida ARERA sul digitale per il settore elettrico.",
    ])

    _add_page_break(doc)


# === Annex C - Regulatory references =====================================


def build_annex_c(doc):
    add_heading(doc, "Allegato C · Riferimenti normativi", level=1)

    add_paragraph(
        doc,
        "Riferimenti normativi citati nel documento. Le voci sono "
        "elencate con riferimento ufficiale per consentire la verifica."
    )

    add_table(
        doc,
        headers=["Riferimento", "Titolo", "Ambito"],
        rows=[
            ["Reg. UE 2016/679 (GDPR)",
             "Regolamento generale sulla protezione dei dati",
             "Dati personali, trattamento, residency"],
            ["Dir. UE 2022/2555 (NIS2)",
             "Direttiva sulla sicurezza delle reti e dei sistemi informativi",
             "Cybersecurity per operatori essenziali; recepimento Italia D.Lgs. 138/2024"],
            ["Reg. UE 2024/1689 (AI Act)",
             "Regolamento sulla intelligenza artificiale",
             "Sistemi AI ad alto rischio, gestione, trasparenza, supervisione umana"],
            ["ISO/IEC 27001:2022",
             "Information security management systems - Requirements",
             "ISMS aziendale (baseline)"],
            ["ISO/IEC 27019:2024",
             "Information security controls for the energy utility industry",
             "Estensione settore energia per ISO 27001"],
            ["ISO/IEC 27017:2015",
             "Code of practice for information security controls based on ISO/IEC 27002 for cloud services",
             "Controlli cloud-specific"],
            ["ISO/IEC 27018:2019",
             "Code of practice for protection of PII in public clouds",
             "Protezione PII in cloud pubblici"],
            ["IEC 62443 serie",
             "Industrial communication networks - Network and system security",
             "Cybersecurity OT/ICS (zone, conduit, security level)"],
            ["IEC 60870-5-104",
             "Telecontrol equipment and systems - Network access for IEC 60870-5-101",
             "Protocollo SCADA su TCP/IP"],
            ["IEC 61850",
             "Communication networks and systems for power utility automation",
             "Standard di comunicazione in substation"],
            ["Linee guida ENISA",
             "Cloud Security guidance for the Energy Sector",
             "Best practice di sicurezza cloud per settori essenziali"],
            ["Direttiva (UE) 2019/944",
             "Norme comuni per il mercato interno dell'energia elettrica",
             "Mercato elettrico UE, customer empowerment, smart metering"],
        ],
        col_widths_cm=[4.5, 6.5, 5.5],
        first_col_bold=True,
    )

    add_paragraph(
        doc,
        "Le normative evolvono. La data di riferimento per le scelte di "
        "questo documento è luglio 2026. Aggiornamenti delle norme (ad "
        "esempio l'adozione progressiva degli atti delegati dell'AI Act) "
        "vanno presidiati con cadenza almeno semestrale dalla funzione "
        "Governance & Compliance."
    )


# === Main ================================================================


def build_document():
    """Costruisce il documento completo e lo restituisce."""
    doc = Document()
    _configure_page(doc)
    _configure_styles(doc)
    _add_header_footer(
        doc,
        "EnergoGrid · Infrastruttura AI multi-cloud · Capstone modulo 05",
    )

    build_cover(doc)
    build_toc(doc)
    build_executive_summary(doc)
    build_section_1(doc)
    build_section_2(doc)
    build_section_3(doc)
    build_section_4(doc)
    build_section_5(doc)
    build_section_6(doc)
    build_section_7(doc)
    build_section_8(doc)
    build_section_9(doc)
    build_section_10(doc)
    build_annex_a(doc)
    build_annex_b(doc)
    build_annex_c(doc)

    return doc


def main():
    doc = build_document()
    out_path = Path(__file__).parent / OUTPUT_FILE
    doc.save(str(out_path))
    print(f"[OK] DOCX scritto in {out_path}")


if __name__ == "__main__":
    main()
